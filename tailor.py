import argparse
import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document

try:
    from llm_generate_resume import LLMResumer  # type: ignore
    _LLM_RESUMER_AVAILABLE = True
except Exception:
    _LLM_RESUMER_AVAILABLE = False

TECH_KEYWORDS = [
    'python','c++','java','javascript','typescript','django','flask','react','reactjs','dash','graphql','rest','api',
    'aws','azure','gcp','lambda','cloudformation','sagemaker','dynamodb','s3','ec2','cloud','databricks','hex',
    'docker','kubernetes','terraform','jenkins','gitlab','github','bitbucket','confluence',
    'mlops','ml','ai','machine','learning','pytorch','tensorflow','keras','scikit-learn','sklearn','pandas','numpy','matplotlib',
    'opencv','snowflake','postgresql','postgres','sqlite','oracle','sql','nosql',
    'microservices','ci/cd','cicd','pipeline','pipelines','monitoring','logging','sonarqube', 'GCP Vertex AI', 'OpenAI GPT-4o'
]

_non_alnum = re.compile(r"[^a-z0-9+#.\-\s]")


def read_text(path: Path) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def tokenize(text: str) -> list[str]:
    text = (text or "").lower()
    text = _non_alnum.sub(" ", text)
    return [t for t in text.split() if len(t) > 1]


def extract_skills(text: str) -> list[str]:
    tokens = set(tokenize(text))
    skills = [k for k in TECH_KEYWORDS if k in tokens]
    # normalize for display
    return sorted({s.upper() for s in skills})


def build_targeted_summary(jd_text: str, matched: list[str]) -> str:
    base = "Results-oriented engineer with experience across MLOps, Full Stack, and Cloud."
    if not matched:
        return base + " Ready to contribute to your team."
    top = ", ".join(matched[:12])
    return f"{base} Aligned to this role with strengths in {top}."


def parse_contact_block(resume_text: str) -> str:
    lines = [l.strip() for l in resume_text.splitlines() if l.strip()]
    return " | ".join(lines[:5])


def write_docx(name: str, contact: str, summary: str, matched: list[str], full_resume: str, out_path: Path) -> None:
    doc = Document()
    doc.add_heading(name or "Candidate", level=0)
    if contact:
        doc.add_paragraph(contact)
    doc.add_paragraph("")

    doc.add_heading("Targeted Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_paragraph("")

    doc.add_heading("Matched Skills", level=1)
    doc.add_paragraph(" · ".join(matched) if matched else "No direct matches detected; general capabilities emphasized.")
    doc.add_paragraph("")

    doc.add_heading("Full Resume", level=1)
    for line in full_resume.splitlines():
        doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    here = Path(__file__).parent
    parser = argparse.ArgumentParser(description="Tailor resume to a job description and generate a .docx.")
    parser.add_argument("--resume", default=str(here.parent / "resume" / "input" / "resume.txt"), help="Path to base resume text file")
    parser.add_argument("--jd", required=True, help="Path to job description text file")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    parser.add_argument("--out", default=str(here / "output" / f"tailored_{ts}.docx"), help="Output .docx path")
    parser.add_argument("--name", default=None, help="Candidate name (fallbacks to first line of resume)")
    parser.add_argument("--company", default="", help="Company name (for LLM tailoring)")
    parser.add_argument("--role", default="", help="Role/title (for LLM tailoring)")
    parser.add_argument("--use-llm", action="store_true", help="Use LLM-based tailoring (requires OPENAI_API_KEY)")
    args = parser.parse_args()

    resume_text = read_text(Path(args.resume))
    jd_text = read_text(Path(args.jd))

    name = args.name or (resume_text.splitlines()[0].strip() if resume_text.splitlines() else "Candidate")
    company = args.company or "Company"
    role = args.role or "Role"

    # Try LLM-based tailoring if requested and available
    if args.use_llm and _LLM_RESUMER_AVAILABLE:
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            try:
                print("[tailor] Using LLM-based resume tailoring...")
                resumer = LLMResumer(openai_api_key=api_key)
                resumer.set_resume_data(resume_text)
                tailored_text = resumer.generate_tailored_resume(jd_text, company, role)
                
                # Create simple document from tailored text
                doc = Document()
                doc.add_heading(name, level=0)
                for line in tailored_text.splitlines():
                    if line.strip():
                        doc.add_paragraph(line.strip())
                
                Path(args.out).parent.mkdir(parents=True, exist_ok=True)
                doc.save(args.out)
                print("LLM-tailored resume generated at:", os.path.abspath(args.out))
                return
            except Exception as e:
                print(f"[tailor] LLM tailoring failed: {e}. Falling back to keyword-based method.")

    # Fallback: keyword-based tailoring
    resume_skills = extract_skills(resume_text)
    jd_skills = extract_skills(jd_text)
    matched = [s for s in jd_skills if s in resume_skills]

    contact = parse_contact_block(resume_text)
    summary = build_targeted_summary(jd_text, matched)

    write_docx(name, contact, summary, matched, resume_text, Path(args.out))
    print("Tailored resume generated at:", os.path.abspath(args.out))


if __name__ == "__main__":
    main()
