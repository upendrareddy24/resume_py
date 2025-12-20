import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

try:
    import openai_compat  # noqa: F401
except Exception:
    openai_compat = None

from docx import Document

from resume_builder_templates import (
    prompt_header_template,
    prompt_education_template,
    prompt_working_experience_template,
    prompt_projects_template,
    prompt_achievements_template,
    prompt_certifications_template,
    prompt_additional_skills_template,
)

try:
    from openai import OpenAI  # type: ignore
    _OPENAI = True
except Exception:
    _OPENAI = False


def _read_cfg(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _build_prompt_block(title: str, base: str, payload: str) -> str:
    return base.format(
        personal_information=payload,
        education_details=payload,
        experience_details=payload,
        projects=payload,
        achievements=payload,
        certifications=payload,
        languages=payload,
        interests=payload,
        skills=payload,
    )


def _call_openai(client: Any, model: str, prompt: str) -> str:
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You write ATS-friendly resume sections, concise and metric-driven."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
        max_tokens=600,
    )
    return (resp.choices[0].message.content or "").strip()


def build_resume_doc(sections: Dict[str, str]) -> Document:
    doc = Document()
    for title, content in sections.items():
        if not content:
            continue
        doc.add_heading(title, level=1)
        for line in content.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
    return doc


def tailor_resume_for_job(
    resume_text: str,
    jd_text: str,
    company: str,
    role: str,
    model: str,
    api_key: str | None,
) -> str:
    """
    Returns tailored resume text (plain text) optimized for the given JD using OpenAI.
    Prefers LLMResumer if available, falls back to direct OpenAI client.
    """
    key = api_key or os.getenv("OPENAI_API_KEY")
    if not key:
        return resume_text
    
    # Try using LLMResumer first (preferred method)
    try:
        from llm_generate_resume import LLMResumer  # type: ignore
        resumer = LLMResumer(openai_api_key=key)
        resumer.set_resume_data(resume_text)
        return resumer.generate_tailored_resume(jd_text, company, role)
    except Exception as e:
        print(f"[resume_builder] LLMResumer unavailable ({e}), using fallback")
    
    # Fallback to direct OpenAI client
    if not _OPENAI:
        return resume_text
    try:
        client = OpenAI(api_key=key)
        system = (
            "You are an expert resume writer specializing in ATS optimization. "
            "Given a resume and job description, rewrite the resume to highlight relevant skills, "
            "experience, and keywords that align with the job requirements while maintaining truthfulness."
        )
        user = (
            f"Company: {company}\nRole: {role}\n\n"
            f"Job Description:\n{jd_text}\n\n"
            f"Current Resume:\n{resume_text}\n\n"
            "Task: Rewrite the resume to emphasize skills and experience matching the job description. "
            "Keep the same structure and facts; optimize keyword placement for ATS. Output plain text only."
        )
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
            temperature=0.5,
            max_tokens=6000,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        print(f"[resume_builder] tailor error: {e}")
        return resume_text


def build_tailored_resume_doc(tailored_text: str) -> Document:
    """
    Creates a simple docx from tailored plain text resume.
    """
    doc = Document()
    for line in tailored_text.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    return doc


def main() -> None:
    cfg_path = Path("config.json")
    if not cfg_path.exists():
        raise SystemExit("config.json not found")
    cfg = _read_cfg(cfg_path)
    rb = cfg.get("resume_builder") or {}
    if not rb:
        raise SystemExit("resume_builder block missing in config.json")

    openai_cfg = cfg.get("openai") or {}
    enabled = bool(openai_cfg.get("enabled"))
    model = (openai_cfg.get("model") or "gpt-4o-mini").strip()
    key = (openai_cfg.get("api_key") or os.getenv("OPENAI_API_KEY") or "").strip()
    if enabled and (not _OPENAI or not key):
        print("OpenAI not available or API key missing; falling back to simple passthrough of inputs.")
        enabled = False

    client = OpenAI(api_key=key) if enabled else None

    def gen(title: str, base_prompt: str, payload_key: str) -> str:
        payload = rb.get(payload_key) or ""
        if not payload:
            return ""
        full_prompt = _build_prompt_block(title, base_prompt, payload)
        if client:
            try:
                return _call_openai(client, model, full_prompt)
            except Exception as e:
                print(f"openai error on {title}:", e)
        return payload

    # Provided prompt strings
    header_base = (
        """
Act as an HR expert and resume writer specializing in ATS-friendly resumes. Your task is to create a professional and polished header for the resume. The header should:

1. **Contact Information**: Include your full name, city and country, phone number, email address, LinkedIn profile, and GitHub profile. Exclude any information that is not provided.
2. **Formatting**: Ensure the contact details are presented clearly and are easy to read.

- **My information:**  
  {personal_information}
""" + prompt_header_template
    )

    education_base = (
        """
Act as an HR expert and resume writer with a specialization in creating ATS-friendly resumes. Your task is to articulate the educational background for a resume. For each educational entry, ensure you include:

1. **Institution Name and Location**: Specify the university or educational institution’s name and location.
2. **Degree and Field of Study**: Clearly indicate the degree earned and the field of study.
3. **Grade**: Include your Grade if it is strong and relevant.
4. **Relevant Coursework**: List key courses with their grades to showcase your academic strengths.

- **My information:**  
  {education_details}
""" + prompt_education_template
    )

    experience_base = (
        """
Act as an HR expert and resume writer with a specialization in creating ATS-friendly resumes. Your task is to detail the work experience for a resume. For each job entry, ensure you include:

1. **Company Name and Location**: Provide the name of the company and its location.
2. **Job Title**: Clearly state your job title.
3. **Dates of Employment**: Include the start and end dates of your employment.
4. **Responsibilities and Achievements**: Describe your key responsibilities and notable achievements, emphasizing measurable results and specific contributions.

- **My information:**  
  {experience_details}
""" + prompt_working_experience_template
    )

    projects_base = (
        """
Act as an HR expert and resume writer with a specialization in creating ATS-friendly resumes. Your task is to highlight notable side projects. For each project, ensure you include:

1. **Project Name and Link**: Provide the name of the project and include a link to the GitHub repository or project page.
2. **Project Details**: Describe any notable recognition or achievements related to the project, such as GitHub stars or community feedback.
3. **Technical Contributions**: Highlight your specific contributions and the technologies used in the project. 

- **My information:**  
  {projects}
""" + prompt_projects_template
    )

    achievements_base = (
        """
Act as an HR expert and resume writer with a specialization in creating ATS-friendly resumes. Your task is to list significant achievements. For each achievement, ensure you include:

1. **Award or Recognition**: Clearly state the name of the award, recognition, scholarship, or honor.
2. **Description**: Provide a brief description of the achievement and its relevance to your career or academic journey.

- **My information:**  
  {achievements}
""" + prompt_achievements_template
    )

    certifications_base = (
        """
Act as an HR expert and resume writer with a specialization in creating ATS-friendly resumes. Your task is to list significant certifications based on the provided details. For each certification, ensure you include:

1. **Certification Name**: Clearly state the name of the certification.
2. **Description**: Provide a brief description of the certification and its relevance to your professional or academic career.

Ensure that the certifications are clearly presented and effectively highlight your qualifications.

If any of the certification details are not provided, omit those sections when filling out the template.

- **My information:**  
  {certifications}
""" + prompt_certifications_template
    )

    addl_skills_base = (
        """
Act as an HR expert and resume writer with a specialization in creating ATS-friendly resumes. Your task is to list additional skills relevant to the job. For each skill, ensure you include:

1. **Skill Category**: Clearly state the category or type of skill.
2. **Specific Skills**: List the specific skills or technologies within each category.
3. **Proficiency and Experience**: Briefly describe your experience and proficiency level.

- **My information:**  
  {languages}
  {interests}
  {skills}
""" + prompt_additional_skills_template
    )

    sections = {
        "Header": gen("Header", header_base, "personal_information"),
        "Education": gen("Education", education_base, "education_details"),
        "Experience": gen("Experience", experience_base, "experience_details"),
        "Projects": gen("Projects", projects_base, "projects"),
        "Achievements": gen("Achievements", achievements_base, "achievements"),
        "Certifications": gen("Certifications", certifications_base, "certifications"),
        "Additional Skills": gen("Additional Skills", addl_skills_base, "additional_skills"),
    }

    # Create document with highlighted skills
    doc = build_resume_doc(sections)
    out = rb.get("out") or f"resume_enhanced_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("Enhanced resume generated:", os.path.abspath(out))


if __name__ == "__main__":
    main()


