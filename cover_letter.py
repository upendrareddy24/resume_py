import argparse
import os
import re
from pathlib import Path
from datetime import datetime

from docx import Document
try:
    from openai import OpenAI  # type: ignore
    _OPENAI_AVAILABLE = True
except Exception:
    _OPENAI_AVAILABLE = False

try:
    # Prefer local LLM adapter when available
    from llm_cover_letter_adapter import LLMCoverLetterJobDescription  # type: ignore
    _LLM_ADAPTER_AVAILABLE = True
except Exception:
    _LLM_ADAPTER_AVAILABLE = False

_non_alnum = re.compile(r"[^a-z0-9+#.\-\s]")


def _tokenize(text: str) -> list[str]:
    text = (text or "").lower()
    text = _non_alnum.sub(" ", text)
    return [t for t in text.split() if len(t) > 1]


CORE_TERMS = [
    "python","java","c++","c#","javascript","typescript","go","sql","nosql",
    "django","flask","fastapi","react","node","graphql","rest","api",
    "ml","ai","machine","learning","deep","pytorch","tensorflow","keras","sklearn","scikit",
    "data","engineer","scientist","analytics","pipeline","etl","spark","airflow","dbt",
    "aws","azure","gcp","lambda","sagemaker","cloudformation","dynamodb","s3","ec2",
    "docker","kubernetes","terraform","jenkins","ansible","gitlab","github",
]


class CoverLetterBuilder:
    def __init__(self, resume_text: str, candidate_name: str = "") -> None:
        self.resume_text = resume_text
        self.candidate_name = candidate_name or "Candidate"

    def extract_keywords(self, jd_text: str, max_terms: int = 24) -> list[str]:
        tokens = set(_tokenize(self.resume_text) + _tokenize(jd_text))
        ordered = [k for k in CORE_TERMS if k in tokens]
        if not ordered:
            ordered = list(tokens)
        out: list[str] = []
        seen = set()
        for t in ordered:
            if t in seen:
                continue
            seen.add(t)
            out.append(t)
            if len(out) >= max_terms:
                break
        return out

    def compute_ats_score(self, jd_text: str) -> int:
        rset = set(_tokenize(self.resume_text))
        jset = set(_tokenize(jd_text))
        if not jset:
            return 0
        overlap = len(rset.intersection(jset))
        score = int(min(100, round(100 * overlap / max(1, len(jset)))))
        if score >= 70 and len(rset.intersection(set(CORE_TERMS))) >= 8:
            score = max(score, 90)
        return score

    def build_docx(self, company: str, role: str, jd_text: str) -> Document:
        doc = Document()
        today = datetime.now().strftime("%B %d, %Y")
        ats = self.compute_ats_score(jd_text)
        top_keywords = ", ".join(self.extract_keywords(jd_text, max_terms=12))

        doc.add_paragraph(today)
        doc.add_paragraph("")
        doc.add_paragraph("Hiring Team")
        if company:
            doc.add_paragraph(company)
        doc.add_paragraph("")

        greet = f"Dear {company} Hiring Team," if company else "Dear Hiring Manager,"
        doc.add_paragraph(greet)

        opening = (
            f"I am excited to apply for the {role or 'role'} at {company or 'your organization'}. "
            f"With hands-on experience across ML/AI, data platforms, and cloud-native engineering, "
            f"I believe my background aligns strongly with your requirements."
        )
        doc.add_paragraph(opening)

        doc.add_paragraph("Highlights aligned to your needs: " + top_keywords)

        p = doc.add_paragraph(); p.style = doc.styles['List Bullet']
        p.add_run("End-to-end ML/AI pipelines (training, deployment, monitoring) on AWS/Azure.")
        p = doc.add_paragraph(); p.style = doc.styles['List Bullet']
        p.add_run("Production APIs and data services (Python, Django/Flask/FastAPI; REST/GraphQL).")
        p = doc.add_paragraph(); p.style = doc.styles['List Bullet']
        p.add_run("DevOps and platform engineering (Docker, Kubernetes, Terraform, CI/CD).")

        if jd_text.strip():
            doc.add_paragraph("")
            doc.add_paragraph("How I address your responsibilities:")
            for line in jd_text.splitlines():
                line = line.strip(" •-\t")
                if len(line) < 4:
                    continue
                li = doc.add_paragraph(); li.style = doc.styles['List Bullet']
                li.add_run(f"{line} – I have relevant experience delivering similar outcomes at scale.")

        doc.add_paragraph("")
        doc.add_paragraph(
            "I am eager to contribute to your team, collaborate cross-functionally, and deliver reliable, high-impact solutions. "
            "Thank you for your time and consideration."
        )
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(self.candidate_name)
        doc.add_paragraph("")
        doc.add_paragraph(f"(Estimated ATS keyword overlap score: ~{ats}%)")
        return doc

    def compose_concise_text(self, jd_text: str, company: str, role: str) -> str:
        rset = set(_tokenize(self.resume_text))
        jset = set(_tokenize(jd_text)) if jd_text else set()
        shared = [t for t in CORE_TERMS if t in rset and (not jset or t in jset)]
        shared = shared[:10] if shared else list(rset)[:10]
        keywords_str = ", ".join(shared)

        resume_lines = [ln.strip() for ln in self.resume_text.splitlines() if ln.strip()]
        example_lines: list[str] = []
        for ln in resume_lines:
            if len(example_lines) >= 3:
                break
            low = ln.lower()
            if any(k in low for k in shared) and 30 <= len(ln) <= 180:
                example_lines.append(ln)
        if not example_lines:
            example_lines = [
                "Built and shipped ML/data services in production with Python and cloud-native tooling.",
                "Owned pipelines, APIs, and CI/CD, improving reliability and developer velocity.",
            ]

        p1 = (
            f"I’m excited about the {role or 'role'} at {company or 'your team'} because the responsibilities align "
            f"with my focus on building reliable, scalable systems. I blend ML/data engineering with strong software "
            f"craft, which maps closely to your requirements around impact and execution."
        )
        p2 = (f"Recent work highlights include: {example_lines[0]} "
              f"{' ' + example_lines[1] if len(example_lines) > 1 else ''} "
              f"{' ' + example_lines[2] if len(example_lines) > 2 else ''} "
              f"Key strengths for this role: {keywords_str}.").strip()
        p3 = (
            f"{company or 'The team'}’s emphasis on thoughtful engineering and real-world outcomes resonates with my approach. "
            f"I’d value the chance to contribute quickly, collaborate across functions, and raise the bar on quality and speed."
        )
        return "\n\n".join([p1, p2, p3])

    def compose_openai_text(self, jd_text: str, company: str, role: str, model: str, api_key: str | None) -> str | None:
        if not _OPENAI_AVAILABLE:
            return None
        try:
            key = api_key or os.getenv("OPENAI_API_KEY")
            if not key:
                return None
            client = OpenAI(api_key=key)
            system = (
                "You are an expert technical recruiter and writing assistant. "
                "Write a concise three-paragraph cover letter without greeting or signature. "
                "Tone: professional, conversational, natural. Use the resume strengths and job requirements."
            )
            user = (
                f"Company: {company}\nRole: {role}\n\n"
                f"Job description:\n{jd_text}\n\n"
                f"Resume:\n{self.resume_text}\n\n"
                "Rules:\n- Three short paragraphs\n- No greeting or signature\n- Reference concrete skills and outcomes "
                "that align with the role\n- Avoid placeholders; write as final text\n"
            )
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                temperature=0.6,
                max_tokens=350,
            )
            return (resp.choices[0].message.content or "").strip()
        except Exception:
            return None


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate a tailored cover letter (prefers local LLM adapter when enabled).")
    ap.add_argument("--config", default="config.json", help="Path to config JSON with cover_letter block")
    args = ap.parse_args()

    cfg_path = Path(args.config)
    if not cfg_path.exists():
        raise SystemExit(f"Config not found: {cfg_path}")
    import json
    cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    cl = cfg.get("cover_letter") or {}
    resume_path = cl.get("resume") or cfg.get("resume")
    jd_path = cl.get("jd")
    name = cl.get("name", "")
    company = cl.get("company", "")
    role = cl.get("role", "")
    out = cl.get("out") or f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    if not resume_path or not jd_path:
        raise SystemExit("cover_letter.resume and cover_letter.jd must be set in config")

    resume_text = _read_text(Path(resume_path))
    jd_text = _read_text(Path(jd_path))
    builder = CoverLetterBuilder(resume_text, name)

    # Prefer LLM adapter when enabled and available
    openai_cfg = cfg.get("openai") or {}
    use_llm = bool(openai_cfg.get("enabled")) and _LLM_ADAPTER_AVAILABLE
    api_key = (openai_cfg.get("api_key") or os.getenv("OPENAI_API_KEY") or "").strip()

    doc: Document
    if use_llm and api_key:
        try:
            # Generate body-only cover letter using local adapter
            adapter = LLMCoverLetterJobDescription(api_key)
            adapter.set_resume(resume_text)
            adapter.set_job_description_from_text(jd_text)
            body_text = adapter.generate_cover_letter().strip()

            # Assemble a clean .docx with greeting/signature around the LLM body
            doc = Document()
            today = datetime.now().strftime("%B %d, %Y")
            doc.add_paragraph(today)
            doc.add_paragraph("")
            doc.add_paragraph("Hiring Team")
            if company:
                doc.add_paragraph(company)
            doc.add_paragraph("")
            greet = f"Dear {company} Hiring Team," if company else "Dear Hiring Manager,"
            doc.add_paragraph(greet)

            for para in [p.strip() for p in body_text.split("\n\n") if p.strip()]:
                doc.add_paragraph(para)

            doc.add_paragraph("")
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph(name or "Candidate")

            # Add ATS estimate footer for context
            ats = builder.compute_ats_score(jd_text)
            doc.add_paragraph("")
            doc.add_paragraph(f"(Estimated ATS keyword overlap score: ~{ats}%)")
        except Exception:
            # Fallback to deterministic builder
            doc = builder.build_docx(company, role, jd_text)
    else:
        # No LLM or missing key → fallback
        doc = builder.build_docx(company, role, jd_text)

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("Cover letter generated:", os.path.abspath(out))


if __name__ == "__main__":
    main()


