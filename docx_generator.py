"""
Word Document (DOCX) Generator for Resumes and Cover Letters
Generates professional documents with precise formatting control
"""
import os
from typing import Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import re

try:
    from llm_experience_parser import parse_experiences_with_llm
    LLM_PARSER_AVAILABLE = True
except ImportError:
    LLM_PARSER_AVAILABLE = False


def _clean_markdown(text: str) -> str:
    """Remove markdown formatting (bold, italic, etc.) from text."""
    if not text:
        return text
    # Convert markdown links [label](url) into "label: url"
    def _link_repl(match: re.Match[str]) -> str:
        label = match.group(1).strip()
        url = match.group(2).strip()
        if not label:
            return url
        if label.lower() in {"github", "linkedin", "portfolio", "website"}:
            return f"{label}: {url}"
        return f"{label} ({url})"

    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', _link_repl, text)
    # Remove bold (**text** or __text__)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    # Remove italic (*text* or _text_) - be careful not to remove list bullets
    text = re.sub(r'(?<!\*)\*(?!\*)([^\*]+)\*(?!\*)', r'\1', text)
    text = re.sub(r'(?<!_)_(?!_)([^_]+)_(?!_)', r'\1', text)
    return text

LINK_REGEX = re.compile(r'https?://[^\s,;]+', re.IGNORECASE)
UNWANTED_PHRASES = [
    "this resume is crafted to align bhavana's extensive technical",
    "this resume is crafted to align bhavana's extensive technical background",
]


def _extract_contact_details(content: str, sections: dict) -> dict[str, str]:
    details: dict[str, str] = {}
    candidates = content.splitlines()[:30]
    contact_section = sections.get("contact")
    if contact_section:
        candidates.extend(contact_section.splitlines())

    email_regex = re.compile(r'[\w\.-]+@[\w\.-]+\.\w+')
    phone_regex = re.compile(r'(\+?\d[\d\s\-\(\)]{7,}\d)')

    for line in candidates:
        if not line:
            continue
        stripped = _clean_markdown(line.strip())
        lower = stripped.lower()
        if "linkedin.com" in lower and "linkedin" not in details:
            match = LINK_REGEX.search(stripped)
            details["linkedin"] = match.group(0) if match else stripped
            continue
        if "github.com" in lower and "github" not in details:
            match = LINK_REGEX.search(stripped)
            details["github"] = match.group(0) if match else stripped
            continue
        match_email = email_regex.search(stripped)
        if match_email and "email" not in details:
            details["email"] = match_email.group(0)
        match_phone = phone_regex.search(stripped)
        if match_phone and "phone" not in details:
            details["phone"] = re.sub(r'\s+', ' ', match_phone.group(0)).strip()
    return details


class WordDocumentGenerator:
    """Generate professional Word documents for resumes and cover letters"""
    
    def __init__(self):
        pass
    
    def _experiences_from_structured(self, work_items: list[dict]) -> list[dict]:
        """Build experiences list from structured YAML work entries."""
        results: list[dict] = []
        if not work_items:
            return results
        for job in work_items:
            try:
                position = (job.get("position") or job.get("title") or "").strip()
                company = (job.get("company") or "").strip()
                start = job.get("startDate") or job.get("start_date") or ""
                end = job.get("endDate") or job.get("end_date") or ""
                dates = " – ".join([p for p in [start, end] if p]).strip()
                location = (job.get("location") or "").strip()
                bullets: list[str] = []
                for key in ("responsibilities", "highlights"):
                    vals = job.get(key) or []
                    if isinstance(vals, list):
                        bullets.extend([str(v).strip() for v in vals if str(v).strip()])
                for key in ("achievements",):
                    vals = job.get(key) or []
                    if isinstance(vals, list):
                        bullets.extend([str(v).strip() for v in vals if str(v).strip()])
                results.append({
                    "position": position,
                    "company": company,
                    "dates": dates,
                    "location": location,
                    "bullets": bullets,
                })
            except Exception:
                continue
        return results
    
    def generate_resume_docx(
        self,
        content: str,
        output_path: str,
        job_title: str = "",
        company_name: str = "",
        candidate_name: str = "",
        structured: dict | None = None
    ) -> bool:
        """
        Generate a professional 3-page resume in DOCX format
        
        Args:
            content: Resume text content
            output_path: Path to save the DOCX
            job_title: Target job title
            company_name: Target company name
            candidate_name: Candidate's name
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Create document
            doc = Document()
            
            # Set margins (1 inch all around)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Parse content into sections
            sections_dict = self._parse_resume_content(content)
            
            print(f"  [docx-debug] Found sections: {list(sections_dict.keys())}")
            
            # Extract or use provided candidate name
            name_to_display = candidate_name
            if not name_to_display and content:
                first_line = content.split('\n')[0].strip()
                if first_line and len(first_line) < 50 and not any(kw in first_line.lower() for kw in ['target', 'position', 'github', '@']):
                    name_to_display = first_line
            
            # Header: Name
            if name_to_display:
                name_para = doc.add_paragraph(name_to_display)
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                name_run = name_para.runs[0]
                name_run.font.size = Pt(18)
                name_run.font.bold = True
                name_run.font.color.rgb = RGBColor(0, 0, 0)
            
            contact_details = _extract_contact_details(content, sections_dict)
            # Display contact info on separate centered lines
            if contact_details.get("email"):
                email_para = doc.add_paragraph(contact_details["email"])
                email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                email_para.paragraph_format.space_after = Pt(3)
                if email_para.runs:
                    email_para.runs[0].font.size = Pt(10)
                    email_para.runs[0].font.color.rgb = RGBColor(80, 80, 80)
            
            if contact_details.get("phone"):
                phone_para = doc.add_paragraph(contact_details["phone"])
                phone_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                phone_para.paragraph_format.space_after = Pt(3)
                if phone_para.runs:
                    phone_para.runs[0].font.size = Pt(10)
                    phone_para.runs[0].font.color.rgb = RGBColor(80, 80, 80)
            
            if contact_details.get("github"):
                github_para = doc.add_paragraph(contact_details["github"])
                github_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                github_para.paragraph_format.space_after = Pt(3)
                if github_para.runs:
                    github_para.runs[0].font.size = Pt(10)
                    github_para.runs[0].font.color.rgb = RGBColor(80, 80, 80)
            
            if contact_details.get("linkedin"):
                linkedin_para = doc.add_paragraph(contact_details["linkedin"])
                linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                linkedin_para.paragraph_format.space_after = Pt(6)
                if linkedin_para.runs:
                    linkedin_para.runs[0].font.size = Pt(10)
                    linkedin_para.runs[0].font.color.rgb = RGBColor(80, 80, 80)

            if contact_line_parts:
                doc.add_paragraph()
            
            # Professional Summary (15 bullet points)
            summary_added = False
            for key in ['summary', 'professional_summary', 'objective']:
                if key in sections_dict and sections_dict[key].strip():
                    # Section header
                    self._add_section_header(doc, "PROFESSIONAL SUMMARY")
                    
                    summary_text = sections_dict[key]
                    summary_lines = [l.strip() for l in summary_text.split('\n') if l.strip()]
                    bullet_count = 0
                    
                    for line in summary_lines:
                        # Skip section headers
                        if line.isupper() and len(line) < 50:
                            continue
                        
                        # Clean up the line
                        clean_line = line.lstrip('•-*►▪→◆ ').strip()
                        
                        if clean_line and len(clean_line) > 20:
                            self._add_bullet_point(doc, clean_line)
                            bullet_count += 1
                            if bullet_count >= 15:
                                break
                    
                    if bullet_count > 0:
                        summary_added = True
                        doc.add_paragraph()  # Spacing
                    break
            
            # Default summary if none found
            if not summary_added:
                self._add_section_header(doc, "PROFESSIONAL SUMMARY")
                self._add_bullet_point(doc, "Experienced professional with a proven track record of delivering high-quality results")
                self._add_bullet_point(doc, "Strong technical skills and ability to work in fast-paced environments")
                self._add_bullet_point(doc, "Excellent communication and collaboration abilities")
                doc.add_paragraph()
            
            # Work Experience
            experience_added = False
            for key in ['experience', 'work_experience', 'employment']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "WORK EXPERIENCE")
                    
                    exp_text = sections_dict[key]
                    # Prefer structured YAML work history if provided
                    experiences = []
                    if structured and isinstance(structured.get("work"), list):
                        experiences = self._experiences_from_structured(structured.get("work"))
                    if not experiences:
                        experiences = self._parse_experiences(exp_text)
                    # Fallbacks: ensure up to 5 distinct experiences by merging from section and whole content
                    def _merge_unique(base: list, extra: list) -> list:
                        seen_local = {(b.get('position',''), b.get('company','')) for b in base}
                        for e in extra:
                            key = (e.get('position',''), e.get('company',''))
                            if key not in seen_local and (key[0] or key[1]):
                                base.append(e)
                                seen_local.add(key)
                        return base
                    if len(experiences) < 5:
                        experiences = _merge_unique(experiences, self._parse_experiences_loose(exp_text))
                    if len(experiences) < 5:
                        experiences = _merge_unique(experiences, self._parse_experiences_loose(content))
                    
                    # Debug: Print parsed experiences
                    print(f"  [docx-debug] Parsed {len(experiences)} experiences")
                    for i, exp in enumerate(experiences[:5]):
                        print(f"  [docx-debug] Exp {i+1}: position='{exp.get('position', 'N/A')}', company='{exp.get('company', 'N/A')}', dates='{exp.get('dates', 'N/A')}', location='{exp.get('location', 'N/A')}'")
                    
                    if experiences:
                        for exp in experiences[:5]:  # Show up to 5 companies
                            # Role and Company
                            position = exp.get('position', 'Position')
                            company = exp.get('company', 'Company')
                            
                            job_para = doc.add_paragraph()
                            job_run = job_para.add_run(f"{position} | {company}")
                            job_run.font.size = Pt(11)
                            job_run.font.bold = True
                            job_para.paragraph_format.space_before = Pt(6)
                            job_para.paragraph_format.space_after = Pt(3)
                            
                            # Date and Location (format: "May 2025 – Present | Remote")
                            if exp.get('dates') or exp.get('location'):
                                date_loc = []
                                if exp.get('dates'):
                                    date_loc.append(exp['dates'])
                                if exp.get('location'):
                                    date_loc.append(exp['location'])
                                
                                date_para = doc.add_paragraph(' | '.join(date_loc))
                                date_run = date_para.runs[0]
                                date_run.font.size = Pt(10)
                                date_run.font.italic = True
                                date_run.font.color.rgb = RGBColor(100, 100, 100)
                                date_para.paragraph_format.space_after = Pt(6)
                            
                            # Bullet points - show all bullets (no limit)
                            for bullet in exp.get('bullets', []):
                                self._add_bullet_point(doc, bullet)
                        
                        experience_added = True
                    else:
                        # Fallback rendering
                        exp_lines = [l.strip() for l in exp_text.split('\n') if l.strip()]
                        for line in exp_lines[:20]:
                            if line.startswith('•') or line.startswith('-'):
                                self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                            else:
                                para = doc.add_paragraph(line)
                                para.paragraph_format.space_after = Pt(6)
                        experience_added = True
                    
                    doc.add_paragraph()  # Spacing
                    break
            
            # Education
            for key in ['education', 'academic']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "EDUCATION")
                    
                    edu_text = sections_dict[key]
                    edu_lines = [l.strip() for l in edu_text.split('\n') if l.strip()]
                    
                    for line in edu_lines:
                        if line.startswith('•') or line.startswith('-'):
                            self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                        else:
                            para = doc.add_paragraph(line)
                            para.paragraph_format.space_after = Pt(4)
                    
                    doc.add_paragraph()  # Spacing
                    break
            
            # Technical Skills
            for key in ['skills', 'technical_skills', 'competencies']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "TECHNICAL SKILLS")
                    
                    skills_text = sections_dict[key].replace('---', '').strip()
                    if skills_text:
                        # Parse skills into bullet points
                        skills_lines = [l.strip() for l in skills_text.split('\n') if l.strip()]
                        for line in skills_lines:
                            # Check if line already has bullet or is a category
                            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                                self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                            elif ':' in line and len(line) < 200:
                                # Category line like "Programming Languages: Python, C++"
                                self._add_bullet_point(doc, line)
                            else:
                                # Plain text - add bullet
                                self._add_bullet_point(doc, line)
                    
                    doc.add_paragraph()  # Spacing
                    break
            
            # Functional Expertise
            for key in ['functional_expertise', 'functional_skills', 'domain_expertise']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "FUNCTIONAL EXPERTISE")
                    
                    expertise_text = sections_dict[key].replace('---', '').strip()
                    if expertise_text:
                        # Parse expertise into bullet points
                        expertise_lines = [l.strip() for l in expertise_text.split('\n') if l.strip()]
                        for line in expertise_lines:
                            # Check if line already has bullet or is a category
                            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                                self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                            elif ':' in line and len(line) < 200:
                                # Category line like "Machine Learning & AI: ML pipeline development"
                                self._add_bullet_point(doc, line)
                            else:
                                # Plain text - add bullet
                                self._add_bullet_point(doc, line)
                    
                    doc.add_paragraph()  # Spacing
                    break
            
            # Key Achievements
            for key in ['achievements', 'key_achievements', 'accomplishments']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "KEY ACHIEVEMENTS")
                    
                    achievement_text = sections_dict[key]
                    achievement_lines = [l.strip() for l in achievement_text.split('\n') if l.strip()]
                    
                    for line in achievement_lines[:10]:
                        if line.startswith('•') or line.startswith('-'):
                            self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                        else:
                            para = doc.add_paragraph(line)
                            para.paragraph_format.space_after = Pt(4)
                    
                    doc.add_paragraph()  # Spacing
                    break
            
            # Publications
            for key in ['publications', 'papers', 'research']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "PUBLICATIONS")
                    
                    pub_text = sections_dict[key]
                    pub_lines = [l.strip() for l in pub_text.split('\n') if l.strip()]
                    
                    for line in pub_lines:
                        if line.startswith('•') or line.startswith('-'):
                            self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                        else:
                            para = doc.add_paragraph(line)
                            para.paragraph_format.space_after = Pt(4)
                    
                    doc.add_paragraph()  # Spacing
                    break
            
            # Projects
            for key in ['projects', 'key_projects']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "KEY PROJECTS")
                    
                    proj_text = sections_dict[key]
                    proj_lines = [l.strip() for l in proj_text.split('\n') if l.strip()]
                    
                    for line in proj_lines[:10]:
                        if line.startswith('•') or line.startswith('-'):
                            self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                        else:
                            para = doc.add_paragraph(line)
                            para.paragraph_format.space_after = Pt(4)
                    
                    doc.add_paragraph()  # Spacing
                    break
            
            # Certifications
            for key in ['certifications', 'certificates']:
                if key in sections_dict and sections_dict[key].strip():
                    self._add_section_header(doc, "CERTIFICATIONS")
                    
                    cert_text = sections_dict[key]
                    cert_lines = [l.strip() for l in cert_text.split('\n') if l.strip()]
                    
                    for line in cert_lines:
                        if line.startswith('•') or line.startswith('-'):
                            self._add_bullet_point(doc, line.lstrip('•-*►▪→◆ '))
                        else:
                            para = doc.add_paragraph(line)
                            para.paragraph_format.space_after = Pt(4)
                    
                    break
            
            # Save document
            doc.save(output_path)
            print(f"[docx] ✅ Resume DOCX generated: {output_path}")
            return True
            
        except Exception as e:
            print(f"[docx] ❌ Error generating resume DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _add_section_header(self, doc, text: str):
        """Add a formatted section header"""
        para = doc.add_paragraph(text)
        run = para.runs[0]
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(8)
        # Add bottom border
        para.paragraph_format.left_indent = Pt(0)
    
    def _add_bullet_point(self, doc, text: str):
        """Add a formatted bullet point"""
        # Clean markdown formatting
        text = _clean_markdown(text)
        para = doc.add_paragraph(text, style='List Bullet')
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.runs[0] if para.runs else None
        if run:
            run.font.size = Pt(10)
    
    def _parse_resume_content(self, content: str) -> dict:
        """Parse resume content into sections"""
        sections = {}
        current_section = None
        current_content = []
        
        lines = content.split('\n')
        
        section_keywords = {
            'contact': ['contact', 'email:', 'phone:', 'github:', 'linkedin:'],
            'summary': ['summary', 'objective', 'profile', 'professional summary'],
            'experience': ['experience', 'employment', 'work history', 'work experience', 'professional experience'],
            'education': ['education', 'academic'],
            'skills': ['skills', 'competencies', 'technical skills'],
            'functional_expertise': ['functional expertise', 'functional skills', 'domain expertise'],
            'achievements': ['key achievements', 'achievements', 'accomplishments'],
            'publications': ['publications', 'papers', 'research'],
            'projects': ['projects', 'portfolio', 'key projects'],
            'certifications': ['certifications', 'licenses', 'certificates']
        }
        
        # Patterns to skip (optimization text, meta descriptions, etc.)
        skip_patterns = [
            r'specifically optimized for',
            r'highlighting.*relevant skills',
            r'this resume (is|has been).*optimized',
            r'this resume (is|has been).*tailored',
            r'this resume (is|has been).*designed',
            r'this resume (is|has been).*crafted',
            r'tailored.*for.*position',
            r'optimized for.*position at',
            r'formatted to align with',
            r'designed to showcase',
            r'this resume is crafted to align',
        ]
        
        for line in lines:
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            if not current_section and not line_stripped:
                continue
            
            # Skip horizontal separator lines (---, ===, etc.)
            if line_stripped and all(c in '-=_*~' for c in line_stripped):
                continue
            
            # Remove markdown bold syntax (**text**)
            line_stripped = re.sub(r'\*\*(.*?)\*\*', r'\1', line_stripped)
            
            # Skip lines matching optimization text patterns
            if any(re.search(pattern, line_lower) for pattern in skip_patterns):
                continue
            
            # Check for section headers
            is_section_header = False
            found_section = None
            
            # All-caps headers
            if line_stripped.isupper() and len(line_stripped) > 3:
                for section_key, keywords in section_keywords.items():
                    if any(keyword in line_lower for keyword in keywords):
                        is_section_header = True
                        found_section = section_key
                        break
            
            # Keyword headers
            if not is_section_header:
                for section_key, keywords in section_keywords.items():
                    if any(line_lower.startswith(keyword) or line_lower == keyword for keyword in keywords):
                        is_section_header = True
                        found_section = section_key
                        break
            
            if is_section_header and found_section:
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                current_section = found_section
                current_content = []
            elif current_section:
                current_content.append(line)
            else:
                # Contact info detection
                if 'github:' in line_lower or 'linkedin:' in line_lower or '@' in line:
                    if 'contact' not in sections:
                        sections['contact'] = line_stripped
                    else:
                        sections['contact'] += '\n' + line_stripped
        
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections

    def generate_cover_letter_docx(
        self,
        content: str,
        output_path: str,
        job_title: str = "",
        company_name: str = "",
        candidate_name: str = "",
        candidate_email: str = "",
        candidate_phone: str = ""
    ) -> bool:
        """
        Generate a simple, professional 1-page cover letter DOCX.
        This is intentionally much simpler than the resume generator.
        """
        try:
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)

            doc = Document()

            # Set basic margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Header: candidate name
            if candidate_name:
                name_para = doc.add_paragraph(candidate_name)
                name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                name_run = name_para.runs[0]
                name_run.font.size = Pt(14)
                name_run.font.bold = True

            # Contact info
            if candidate_email or candidate_phone:
                contact_parts = []
                if candidate_email:
                    contact_parts.append(candidate_email)
                if candidate_phone:
                    contact_parts.append(candidate_phone)
                contact_para = doc.add_paragraph(" | ".join(contact_parts))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if contact_para.runs:
                    contact_para.runs[0].font.size = Pt(10)

            # Add a blank line
            doc.add_paragraph()

            # Date (optional)
            # We could add today's date here if needed

            # Company + role line (optional)
            if company_name or job_title:
                target_line = ""
                if company_name and job_title:
                    target_line = f"{company_name} – {job_title}"
                elif company_name:
                    target_line = company_name
                elif job_title:
                    target_line = job_title
                tgt_para = doc.add_paragraph(target_line)
                tgt_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if tgt_para.runs:
                    tgt_para.runs[0].font.size = Pt(11)

                doc.add_paragraph()

            # Main body content (from LLM or fallback template)
            if content:
                for line in content.split("\n"):
                    doc.add_paragraph(line)

            # Save document
            doc.save(output_path)
            print(f"[docx] ✅ Cover letter DOCX generated: {output_path}")
            return True
        except Exception as e:
            print(f"[docx] ❌ Error generating cover letter DOCX: {e}")
            return False
    
    def _parse_experiences(self, text: str) -> list:
        """Parse work experience into structured format - uses LLM when available"""
        # Try LLM parsing first if available
        if LLM_PARSER_AVAILABLE:
            try:
                print("  [docx-debug] Attempting LLM-based experience parsing...")
                llm_experiences = parse_experiences_with_llm(text)
                if llm_experiences:
                    print(f"  [docx-debug] LLM parsed {len(llm_experiences)} experiences successfully")
                    # Debug output
                    for idx, exp in enumerate(llm_experiences, 1):
                        print(f"    [docx-debug] Exp {idx}: position={exp.get('position')!r}, company={exp.get('company')!r}, dates={exp.get('dates')!r}, location={exp.get('location')!r}")
                    return llm_experiences
                else:
                    print("  [docx-debug] LLM parsing returned no results, falling back to regex parser")
            except Exception as e:
                print(f"  [docx-debug] LLM parsing failed: {e}, falling back to regex parser")
        
        # Fallback to original regex-based parsing
        print("  [docx-debug] Using regex-based experience parsing...")
        experiences = []
        current_exp = None
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Company/position line
            if '|' in line or any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'lead', 'architect', 'scientist', 'consultant', 'owner', 'master', 'scrum']):
                # Save previous experience (even if no bullets, as long as we have position/company)
                if current_exp and (current_exp.get('position') or current_exp.get('company')):
                    experiences.append(current_exp)
                
                current_exp = {'bullets': []}
                
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')]
                    # Company is always the LAST part, position is everything before it
                    if len(parts) >= 2:
                        current_exp['position'] = ' | '.join(parts[:-1])  # All parts except last
                        current_exp['company'] = parts[-1]  # Last part is company
                    elif len(parts) == 1:
                        current_exp['position'] = parts[0]
                else:
                    current_exp['position'] = line
            
            # Date/location line
            elif current_exp and (any(month in line for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec', '20', 'Present']) or '–' in line or '-' in line):
                if '|' in line:
                    parts = line.split('|')
                    current_exp['dates'] = parts[0].strip()
                    if len(parts) > 1:
                        current_exp['location'] = parts[1].strip()
                else:
                    # Check if it's a date range (contains – or -)
                    if '–' in line or ('-' in line and any(month in line for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'])):
                        current_exp['dates'] = line
                    elif any(month in line for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']):
                        current_exp['dates'] = line
                    else:
                        current_exp['location'] = line
            
            # Bullet point
            elif current_exp:
                bullet = line.lstrip('•-*►▪→◆ ')
                if bullet and len(bullet) > 15:
                    current_exp['bullets'].append(bullet)
        
        # Save last experience (even if no bullets, as long as we have position/company)
        if current_exp and (current_exp.get('position') or current_exp.get('company')):
            experiences.append(current_exp)
        
        return experiences

    def _parse_experiences_loose(self, text: str) -> list:
        """Looser parser to capture position|company lines and attach following details."""
        exps = []
        cur = None
        for line in text.split('\n'):
            s = line.strip()
            if not s:
                continue
            if '|' in s:
                parts = [p.strip() for p in s.split('|')]
                if len(parts) >= 2:
                    if cur and (cur.get('position') or cur.get('company')):
                        exps.append(cur)
                    cur = {'position': ' | '.join(parts[:-1]), 'company': parts[-1], 'bullets': []}
                    continue
            if cur and (any(m in s for m in ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec','Present']) or '–' in s or '-' in s):
                if '|' in s:
                    p = [p.strip() for p in s.split('|')]
                    cur['dates'] = p[0]
                    if len(p) > 1:
                        cur['location'] = p[1]
                else:
                    cur['dates'] = s
                continue
            if cur:
                b = s.lstrip('•-*►▪→◆ ')
                if b:
                    cur.setdefault('bullets', []).append(b)
        if cur and (cur.get('position') or cur.get('company')):
            exps.append(cur)
        return exps


# Convenience function
def generate_resume_docx(content: str, output_path: str, **kwargs) -> bool:
    """Generate a resume DOCX"""
    generator = WordDocumentGenerator()
    return generator.generate_resume_docx(content, output_path, **kwargs)

